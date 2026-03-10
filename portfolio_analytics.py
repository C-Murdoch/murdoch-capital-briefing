"""
=============================================================================
MURDOCH CAPITAL - PORTFOLIO ANALYTICS ENGINE v2.0
Cameron Murdoch Portfolio Management | £5,000 Moderate Risk | 10-Year Horizon
=============================================================================

SETUP:
    pip install yfinance pandas numpy requests python-docx

RUN:
    python portfolio_analytics.py

OUTPUTS:
    Daily Reports/YYYY-MM-DD/portfolio_dashboard_YYYY-MM-DD.html   Full investor dashboard
    Daily Reports/YYYY-MM-DD/portfolio_report_YYYY-MM-DD.docx      Word report (view in Word/OneDrive)
    Daily Reports/YYYY-MM-DD/portfolio_state.json                   Saved portfolio state

BROKER SWITCH:
    Change DATA_PROVIDER = "yahoo"  ->  DATA_PROVIDER = "ibkr"
    once Interactive Brokers TWS/Gateway is running locally on port 7497.
=============================================================================
"""

import yfinance as yf
import pandas as pd
import numpy as np
import json
import os
from datetime import datetime, date, timedelta
from abc import ABC, abstractmethod
import warnings
warnings.filterwarnings("ignore")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────

DATA_PROVIDER     = "yahoo"          # Switch to "ibkr" once IB account is live
PORTFOLIO_CAPITAL = 5000.00          # GBP
INCEPTION_DATE    = "2026-03-09"
BACKTEST_MONTHS   = 12               # Months of history to simulate returns
RISK_FREE_RATE    = 0.043            # UK base rate proxy (annualised)
BENCHMARK_TICKER  = "SPY"           # S&P 500 for comparison

PORTFOLIO = {
    # ── CORE GROWTH SLEEVE (50%) ──────────────────────────────────────────
    "ASML":   {"name": "ASML Holding NV",        "sleeve": "Core",       "target_pct": 0.07, "sector": "Technology",  "theme": "EUV Lithography Monopoly",            "status": "HELD"},
    "MSFT":   {"name": "Microsoft Corporation",  "sleeve": "Core",       "target_pct": 0.07, "sector": "Technology",  "theme": "Enterprise AI / Cloud (Azure)",        "status": "HELD"},
    "GOOG":   {"name": "Alphabet Inc",           "sleeve": "Core",       "target_pct": 0.06, "sector": "Technology",  "theme": "AI Search / DeepMind / GCP",           "status": "HELD"},
    "TSM":    {"name": "TSMC ADR",               "sleeve": "Core",       "target_pct": 0.06, "sector": "Technology",  "theme": "Advanced Semiconductor Mfg",           "status": "HELD"},
    "TSLA":   {"name": "Tesla Inc",              "sleeve": "Core",       "target_pct": 0.06, "sector": "Consumer",    "theme": "Autonomy / Optimus Robotics / Energy", "status": "HELD"},
    "LLY":    {"name": "Eli Lilly & Co",         "sleeve": "Core",       "target_pct": 0.06, "sector": "Healthcare",  "theme": "GLP-1 Drugs / AI Drug Discovery",      "status": "HELD"},
    "NEE":    {"name": "NextEra Energy Inc",     "sleeve": "Core",       "target_pct": 0.06, "sector": "Utilities",   "theme": "Clean Energy / AI Data Centre Power",  "status": "HELD"},
    "DE":     {"name": "Deere & Company",        "sleeve": "Core",       "target_pct": 0.06, "sector": "Industrials", "theme": "Precision Agriculture / Autonomy",     "status": "HELD"},
    # ── SATELLITE SLEEVE (15%) ────────────────────────────────────────────
    "IONQ":   {"name": "IonQ Inc",               "sleeve": "Satellite",  "target_pct": 0.04, "sector": "Technology",  "theme": "Quantum Computing",                    "status": "HELD"},
    "CCJ":    {"name": "Cameco Corporation",     "sleeve": "Satellite",  "target_pct": 0.04, "sector": "Energy",      "theme": "Uranium / Nuclear Renaissance",         "status": "HELD"},
    "AMD":    {"name": "Advanced Micro Devices", "sleeve": "Satellite",  "target_pct": 0.04, "sector": "Technology",  "theme": "AI Accelerators / Data Center CPU",    "status": "HELD"},
    "PLTR":   {"name": "Palantir Technologies",  "sleeve": "Satellite",  "target_pct": 0.03, "sector": "Technology",  "theme": "AI Analytics / Government Platform",   "status": "HELD"},
    # ── FIXED INCOME (15%) ───────────────────────────────────────────────
    "SHY":    {"name": "iShares 1-3yr Treasury", "sleeve": "FixedIncome","target_pct": 0.08, "sector": "Fixed Income","theme": "Short Duration US Treasuries",          "status": "HELD"},
    "IGLT.L": {"name": "iShares UK Gilts",       "sleeve": "FixedIncome","target_pct": 0.07, "sector": "Fixed Income","theme": "UK Government Bonds",                   "status": "HELD"},
    # ── COMMODITIES (10%) ────────────────────────────────────────────────
    "IAU":    {"name": "iShares Gold Trust",     "sleeve": "Commodities","target_pct": 0.06, "sector": "Commodities", "theme": "Gold / Inflation Hedge",                "status": "HELD"},
    "PHAG.L": {"name": "WisdomTree Silver",      "sleeve": "Commodities","target_pct": 0.02, "sector": "Commodities", "theme": "Silver / Industrial Demand",            "status": "HELD"},
    "SPUT":   {"name": "Sprott Uranium Trust",   "sleeve": "Commodities","target_pct": 0.02, "sector": "Commodities", "theme": "Physical Uranium",                      "status": "HELD"},
    # ── WATCHLIST ─────────────────────────────────────────────────────────
    "ARM":    {"name": "ARM Holdings plc",       "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "Chip IP Royalties / Architecture",      "status": "WATCH"},
    "CRWD":   {"name": "CrowdStrike Holdings",   "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "AI-Native Cybersecurity",               "status": "WATCH"},
    "ISRG":   {"name": "Intuitive Surgical",     "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Healthcare",  "theme": "Surgical Robotics",                     "status": "WATCH"},
    "CRSP":   {"name": "CRISPR Therapeutics",    "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Healthcare",  "theme": "Gene Editing / Precision Medicine",     "status": "WATCH"},
    "AVGO":   {"name": "Broadcom Inc",           "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "Custom AI ASICs / Enterprise Software", "status": "WATCH"},
    "MRVL":   {"name": "Marvell Technology",     "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "Custom AI Silicon / Optical",           "status": "WATCH"},
    "AMZN":   {"name": "Amazon.com Inc",         "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "AWS Cloud / AI / Logistics",            "status": "WATCH"},
    "ALAB":   {"name": "Astera Labs Inc",        "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "AI Data Centre Connectivity",           "status": "WATCH"},
    "NVDA":   {"name": "NVIDIA Corporation",     "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Technology",  "theme": "AI Infrastructure (re-entry monitor)",  "status": "WATCH"},
    "ABBV":   {"name": "AbbVie Inc",             "sleeve": "Watchlist",  "target_pct": 0.00, "sector": "Healthcare",  "theme": "Immunology / Oncology Pipeline",        "status": "WATCH"},
}


# ─────────────────────────────────────────────────────────────────────────────
# DATA PROVIDER ABSTRACTION
# ─────────────────────────────────────────────────────────────────────────────

class DataProvider(ABC):
    @abstractmethod
    def get_history(self, ticker: str, start: str, end: str) -> pd.DataFrame: ...
    @abstractmethod
    def get_fundamentals(self, ticker: str) -> dict: ...
    @abstractmethod
    def get_fx_rate(self, from_ccy: str, to_ccy: str) -> float: ...


class YahooFinanceProvider(DataProvider):
    def get_history(self, ticker, start, end):
        try:
            df = yf.download(ticker, start=start, end=end, progress=False, auto_adjust=True)
            return df
        except Exception:
            return pd.DataFrame()

    def get_fundamentals(self, ticker):
        try:
            t    = yf.Ticker(ticker)
            info = t.info
            next_earnings = "N/A"
            try:
                cal = t.calendar
                if cal is not None:
                    ed = cal.get("Earnings Date", [])
                    if ed:
                        next_earnings = str(ed[0])[:10]
            except Exception:
                pass
            return {
                "market_cap":        info.get("marketCap"),
                "pe_trailing":       info.get("trailingPE"),
                "pe_forward":        info.get("forwardPE"),
                "peg_ratio":         info.get("pegRatio"),
                "ev_ebitda":         info.get("enterpriseToEbitda"),
                "ev_revenue":        info.get("enterpriseToRevenue"),
                "revenue_growth":    info.get("revenueGrowth"),
                "earnings_growth":   info.get("earningsGrowth"),
                "gross_margins":     info.get("grossMargins"),
                "operating_margins": info.get("operatingMargins"),
                "free_cashflow":     info.get("freeCashflow"),
                "debt_to_equity":    info.get("debtToEquity"),
                "return_on_equity":  info.get("returnOnEquity"),
                "analyst_rec":       info.get("recommendationKey", "n/a"),
                "analyst_count":     info.get("numberOfAnalystOpinions"),
                "target_mean":       info.get("targetMeanPrice"),
                "target_high":       info.get("targetHighPrice"),
                "target_low":        info.get("targetLowPrice"),
                "current_price":     info.get("currentPrice") or info.get("regularMarketPrice"),
                "52w_high":          info.get("fiftyTwoWeekHigh"),
                "52w_low":           info.get("fiftyTwoWeekLow"),
                "beta":              info.get("beta"),
                "dividend_yield":    info.get("dividendYield"),
                "shares_short_pct":  info.get("shortPercentOfFloat"),
                "next_earnings":     next_earnings,
                "currency":          info.get("currency", "USD"),
            }
        except Exception as e:
            return {"error": str(e)}

    def get_fx_rate(self, from_ccy, to_ccy):
        try:
            hist = yf.Ticker(f"{from_ccy}{to_ccy}=X").history(period="1d")
            if not hist.empty:
                return float(hist["Close"].iloc[-1])
        except Exception:
            pass
        return 1.0


class InteractiveBrokersProvider(DataProvider):
    """
    Stub for Interactive Brokers TWS API.
    To activate:
      1. pip install ibapi
      2. Run IB Gateway or Trader Workstation locally (port 7497)
      3. Implement the three methods below using the IB API docs:
         https://interactivebrokers.github.io/tws-api/
    """
    def get_history(self, ticker, start, end):
        raise NotImplementedError("Implement reqHistoricalData() via ibapi")

    def get_fundamentals(self, ticker):
        raise NotImplementedError("Implement reqFundamentalData() via ibapi")

    def get_fx_rate(self, from_ccy, to_ccy):
        raise NotImplementedError("Implement reqMktData() for FX via ibapi")


def get_provider() -> DataProvider:
    return InteractiveBrokersProvider() if DATA_PROVIDER == "ibkr" else YahooFinanceProvider()


# ─────────────────────────────────────────────────────────────────────────────
# TECHNICAL INDICATORS
# ─────────────────────────────────────────────────────────────────────────────

def calc_rsi(prices, period=14):
    if len(prices) < period + 1:
        return float("nan")
    d = prices.diff().dropna()
    g = d.where(d > 0, 0).rolling(period).mean()
    l = (-d.where(d < 0, 0)).rolling(period).mean()
    return round(float((100 - 100 / (1 + g / (l + 1e-10))).iloc[-1]), 2)

def calc_macd(prices, fast=12, slow=26, signal=9):
    if len(prices) < slow + signal:
        return {"macd": None, "signal_line": None, "histogram": None, "crossover": None}
    ef = prices.ewm(span=fast, adjust=False).mean()
    es = prices.ewm(span=slow, adjust=False).mean()
    m  = ef - es
    sl = m.ewm(span=signal, adjust=False).mean()
    h  = m - sl
    co = None
    if len(m) >= 2:
        if m.iloc[-2] < sl.iloc[-2] and m.iloc[-1] > sl.iloc[-1]: co = "BULLISH"
        elif m.iloc[-2] > sl.iloc[-2] and m.iloc[-1] < sl.iloc[-1]: co = "BEARISH"
    return {"macd": round(float(m.iloc[-1]), 4), "signal_line": round(float(sl.iloc[-1]), 4),
            "histogram": round(float(h.iloc[-1]), 4), "crossover": co}

def calc_bollinger(prices, period=20, sd=2):
    if len(prices) < period:
        return {"upper": None, "middle": None, "lower": None, "pct_b": None, "squeeze": False}
    mean  = prices.rolling(period).mean()
    std   = prices.rolling(period).std()
    upper = mean + sd * std
    lower = mean - sd * std
    curr  = float(prices.iloc[-1])
    bw    = float(upper.iloc[-1]) - float(lower.iloc[-1])
    pct_b = (curr - float(lower.iloc[-1])) / (bw + 1e-10)
    bw_series = (upper - lower) / mean
    avg_bw  = float(bw_series.rolling(period).mean().iloc[-1]) if len(prices) >= period * 2 else float("nan")
    curr_bw = float(bw_series.iloc[-1])
    return {"upper": round(float(upper.iloc[-1]), 2), "middle": round(float(mean.iloc[-1]), 2),
            "lower": round(float(lower.iloc[-1]), 2), "pct_b": round(pct_b, 3),
            "squeeze": (curr_bw < avg_bw * 0.7) if not np.isnan(avg_bw) else False}

def calc_sma(prices, period):
    return round(float(prices.rolling(period).mean().iloc[-1]), 2) if len(prices) >= period else float("nan")

def calc_atr(high, low, close, period=14):
    if len(close) < 2: return float("nan")
    tr = pd.concat([(high - low), (high - close.shift(1)).abs(), (low - close.shift(1)).abs()], axis=1).max(axis=1)
    return round(float(tr.ewm(span=period, adjust=False).mean().iloc[-1]), 4)

def calc_volume(volume):
    if len(volume) < 20: return {"ratio": None, "spike": False}
    avg  = float(volume.rolling(20).mean().iloc[-1])
    curr = float(volume.iloc[-1])
    r    = curr / (avg + 1e-10)
    return {"avg_20d": int(avg), "current": int(curr), "ratio": round(r, 2), "spike": r > 1.5}


# ─────────────────────────────────────────────────────────────────────────────
# PERFORMANCE METRICS
# ─────────────────────────────────────────────────────────────────────────────

def calc_metrics(prices: pd.Series, rf=RISK_FREE_RATE) -> dict:
    if len(prices) < 10: return {}
    ret  = prices.pct_change().dropna()
    n    = len(ret)
    trad = 252
    total_ret = (prices.iloc[-1] / prices.iloc[0]) - 1
    n_yr      = n / trad
    cagr      = (1 + total_ret) ** (1 / n_yr) - 1 if n_yr > 0 else 0
    vol       = float(ret.std()) * np.sqrt(trad)
    daily_rf  = rf / trad
    excess    = ret - daily_rf
    sharpe    = float(excess.mean() / (excess.std() + 1e-10)) * np.sqrt(trad)
    down      = ret[ret < daily_rf]
    d_std     = float(down.std()) * np.sqrt(trad) if len(down) > 0 else 1e-10
    sortino   = (cagr - rf) / (d_std + 1e-10)
    cum       = (1 + ret).cumprod()
    max_dd    = float(((cum - cum.cummax()) / cum.cummax()).min())
    return {
        "total_return_pct":  round(total_ret * 100, 2),
        "cagr_pct":          round(cagr * 100, 2),
        "volatility_pct":    round(vol * 100, 2),
        "sharpe_ratio":      round(sharpe, 3),
        "sortino_ratio":     round(sortino, 3),
        "max_drawdown_pct":  round(max_dd * 100, 2),
        "best_day_pct":      round(float(ret.max()) * 100, 2),
        "worst_day_pct":     round(float(ret.min()) * 100, 2),
        "win_rate_pct":      round(float((ret > 0).mean()) * 100, 1),
    }

def calc_beta(pos_ret: pd.Series, bm_ret: pd.Series) -> float:
    aligned = pd.concat([pos_ret, bm_ret], axis=1).dropna()
    if len(aligned) < 10: return float("nan")
    cov = np.cov(aligned.iloc[:, 0], aligned.iloc[:, 1])
    return round(float(cov[0][1] / (cov[1][1] + 1e-10)), 3)


# ─────────────────────────────────────────────────────────────────────────────
# SIGNAL ENGINE
# ─────────────────────────────────────────────────────────────────────────────

def generate_signal(technical: dict, fundamental: dict) -> dict:
    score   = 0
    reasons = []

    r = technical.get("rsi", 50)
    if isinstance(r, float) and not np.isnan(r):
        if r < 30:   score += 2; reasons.append(f"RSI oversold ({r:.0f}) - potential entry")
        elif r > 70: score -= 2; reasons.append(f"RSI overbought ({r:.0f}) - consider trimming")
        elif 40 <= r <= 60: score += 1

    m = technical.get("macd", {})
    if m.get("crossover") == "BULLISH":  score += 2; reasons.append("MACD bullish crossover")
    elif m.get("crossover") == "BEARISH": score -= 2; reasons.append("MACD bearish crossover")
    if m.get("histogram") is not None:
        score += (1 if m["histogram"] > 0 else -1)

    bb = technical.get("bollinger", {})
    pb = bb.get("pct_b")
    if pb is not None:
        if pb < 0.1:  score += 2; reasons.append("Near Bollinger lower band - mean reversion")
        elif pb > 0.9: score -= 1; reasons.append("Near Bollinger upper band")

    price  = technical.get("price", 0)
    s50    = technical.get("sma50")
    s200   = technical.get("sma200")
    if price and s50 and s200 and not (isinstance(s50, float) and np.isnan(s50)):
        if price > s50 > s200:   score += 2; reasons.append("Above SMA50 and SMA200 - strong uptrend")
        elif price < s50 < s200: score -= 2; reasons.append("Below both SMAs - downtrend in force")
        elif s50 > s200:         score += 1; reasons.append("Golden cross in effect")

    if technical.get("volume", {}).get("spike"):
        reasons.append(f"Volume spike ({technical['volume'].get('ratio', 0):.1f}x avg)")

    fw = fundamental.get("pe_forward")
    if fw:
        try:
            fw = float(fw)
            if fw < 15:   score += 1; reasons.append(f"Forward P/E attractive ({fw:.1f}x)")
            elif fw > 60: score -= 1; reasons.append(f"Forward P/E elevated ({fw:.1f}x)")
        except Exception: pass

    rg = fundamental.get("revenue_growth")
    if rg:
        try:
            rg = float(rg)
            if rg > 0.20:  score += 1; reasons.append(f"Revenue growth {rg*100:.0f}% YoY - strong")
            elif rg < 0:   score -= 1; reasons.append(f"Revenue declining {rg*100:.0f}% YoY")
        except Exception: pass

    ar = (fundamental.get("analyst_rec") or "").lower()
    if ar in ["strong_buy", "buy"]:   score += 1; reasons.append(f"Consensus: {ar.replace('_',' ').upper()}")
    elif ar in ["sell", "strong_sell"]: score -= 1; reasons.append(f"Consensus: {ar.replace('_',' ').upper()}")

    target = fundamental.get("target_mean")
    curr   = fundamental.get("current_price") or price
    if target and curr:
        try:
            up = (float(target) - float(curr)) / float(curr)
            if up > 0.20:    score += 1; reasons.append(f"Analyst target +{up*100:.0f}% upside")
            elif up < -0.10: score -= 1; reasons.append(f"Analyst target {up*100:.0f}% downside")
        except Exception: pass

    if   score >= 5:  sig, col = "STRONG BUY",  "#1a7a4a"
    elif score >= 2:  sig, col = "BUY",          "#2e6ba8"
    elif score <= -4: sig, col = "SELL",         "#a31515"
    elif score <= -2: sig, col = "REDUCE",       "#e8a020"
    else:             sig, col = "HOLD",         "#555555"

    return {"signal": sig, "color": col, "score": score, "reasons": reasons[:4]}


# ─────────────────────────────────────────────────────────────────────────────
# PORTFOLIO SIMULATION (BACKTEST)
# ─────────────────────────────────────────────────────────────────────────────

def simulate_portfolio(provider: DataProvider, months_back: int = 12) -> dict:
    end_date   = date.today()
    start_date = end_date - timedelta(days=months_back * 31)
    print(f"\n  Simulating {months_back}-month backtest ({start_date} to {end_date})...")

    held = [t for t, d in PORTFOLIO.items() if d["status"] == "HELD" and d["target_pct"] > 0]
    all_prices = {}

    for ticker in held + [BENCHMARK_TICKER]:
        h = provider.get_history(ticker, str(start_date), str(end_date))
        if not h.empty:
            c = h["Close"].squeeze()
            if isinstance(c, pd.DataFrame): c = c.iloc[:, 0]
            all_prices[ticker] = c

    if len(all_prices) < 3:
        return {"error": "Insufficient price data"}

    df = pd.DataFrame(all_prices).dropna(how="all").ffill()
    bm_returns = df[BENCHMARK_TICKER].pct_change().dropna() if BENCHMARK_TICKER in df.columns else None
    portfolio_daily = pd.Series(0.0, index=df.index)
    results = {}

    for ticker in held:
        if ticker not in df.columns: continue
        alloc  = PORTFOLIO_CAPITAL * PORTFOLIO[ticker]["target_pct"]
        prices = df[ticker].dropna()
        if len(prices) < 5: continue

        start_p = float(prices.iloc[0])
        shares  = alloc / start_p
        dv      = prices * shares

        m = calc_metrics(prices)
        beta = calc_beta(prices.pct_change().dropna(), bm_returns) if bm_returns is not None else float("nan")

        results[ticker] = {
            "allocation_gbp":  round(alloc, 2),
            "start_price":     round(start_p, 2),
            "end_price":       round(float(prices.iloc[-1]), 2),
            "simulated_value": round(float(dv.iloc[-1]), 2),
            "gain_loss":       round(float(dv.iloc[-1]) - alloc, 2),
            "return_pct":      round(((float(prices.iloc[-1]) / start_p) - 1) * 100, 2),
            "metrics":         m,
            "beta":            beta,
        }
        portfolio_daily = portfolio_daily.add(dv, fill_value=0)

    portfolio_daily = portfolio_daily[portfolio_daily > 0]
    pm = calc_metrics(portfolio_daily)
    cash = PORTFOLIO_CAPITAL * 0.10
    total_sim  = sum(r["simulated_value"] for r in results.values()) + cash
    total_ret  = (total_sim / PORTFOLIO_CAPITAL - 1) * 100

    bm_ret = None
    if BENCHMARK_TICKER in df.columns:
        bp = df[BENCHMARK_TICKER].dropna()
        bm_ret = round(((float(bp.iloc[-1]) / float(bp.iloc[0])) - 1) * 100, 2)

    monthly = portfolio_daily.resample("ME").last().round(2) if not portfolio_daily.empty else pd.Series()

    return {
        "simulation_period":    f"{start_date} to {end_date}",
        "total_capital_gbp":    PORTFOLIO_CAPITAL,
        "total_simulated_gbp":  round(total_sim, 2),
        "total_gain_loss_gbp":  round(total_sim - PORTFOLIO_CAPITAL, 2),
        "total_return_pct":     round(total_ret, 2),
        "benchmark_return_pct": bm_ret,
        "alpha_pct":            round(total_ret - (bm_ret or 0), 2),
        "portfolio_metrics":    pm,
        "positions":            results,
        "monthly_values":       {str(k)[:10]: float(v) for k, v in monthly.items()},
    }


# ─────────────────────────────────────────────────────────────────────────────
# SINGLE TICKER ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

def analyse_ticker(ticker: str, provider: DataProvider, hist: pd.DataFrame = None) -> dict:
    end   = date.today()
    start = end - timedelta(days=400)
    if hist is None or hist.empty:
        hist = provider.get_history(ticker, str(start), str(end))
    if hist.empty or len(hist) < 20:
        return {"error": "Insufficient data", "config": PORTFOLIO.get(ticker, {})}

    def squeeze(s):
        s = s.squeeze()
        return s.iloc[:, 0] if isinstance(s, pd.DataFrame) else s

    close  = squeeze(hist["Close"])
    high   = squeeze(hist["High"])
    low    = squeeze(hist["Low"])
    volume = squeeze(hist["Volume"])
    price  = float(close.iloc[-1])
    prev   = float(close.iloc[-2]) if len(close) > 1 else price

    rets = {}
    for label, days in [("1d",1),("1w",5),("1m",21),("3m",63),("6m",126),("1y",252)]:
        rets[label] = round(((price / float(close.iloc[-days-1])) - 1) * 100, 2) if len(close) > days else None

    p52h = round(((price / float(close.rolling(252).max().iloc[-1])) - 1) * 100, 2) if len(close) >= 252 else None
    p52l = round(((price / float(close.rolling(252).min().iloc[-1])) - 1) * 100, 2) if len(close) >= 252 else None

    technical = {
        "price":              round(price, 2),
        "day_change_pct":     round(((price - prev) / prev) * 100, 2),
        "rsi":                calc_rsi(close),
        "macd":               calc_macd(close),
        "bollinger":          calc_bollinger(close),
        "sma20":              calc_sma(close, 20),
        "sma50":              calc_sma(close, 50),
        "sma200":             calc_sma(close, 200),
        "atr":                calc_atr(high, low, close),
        "volume":             calc_volume(volume),
        "returns":            rets,
        "pct_from_52w_high":  p52h,
        "pct_from_52w_low":   p52l,
    }
    fundamental = provider.get_fundamentals(ticker)
    signal      = generate_signal(technical, fundamental)
    metrics     = calc_metrics(close)

    return {
        "config":      PORTFOLIO.get(ticker, {}),
        "technical":   technical,
        "fundamental": fundamental,
        "signal":      signal,
        "metrics":     metrics,
    }


# ─────────────────────────────────────────────────────────────────────────────
# HTML DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────

def generate_dashboard(analysis: dict, sim: dict, report_date: str) -> str:

    def pc(v, rev=False):
        if v is None or (isinstance(v, float) and np.isnan(v)): return "#888"
        return ("#1a7a4a" if not rev else "#a31515") if float(v) > 0 else ("#a31515" if not rev else "#1a7a4a")

    def fmt(v, pre="", suf="", dec=2, sign=False):
        if v is None or (isinstance(v, float) and np.isnan(v)): return "<span style='color:#aaa'>N/A</span>"
        try: v = float(v)
        except Exception: return str(v)
        s = f"+{v:.{dec}f}" if sign and v > 0 else f"{v:.{dec}f}"
        return f"{pre}{s}{suf}"

    def fmt_mc(v):
        if not v: return "N/A"
        try:
            v = float(v)
            if v >= 1e12: return f"${v/1e12:.2f}T"
            if v >= 1e9:  return f"${v/1e9:.1f}B"
            if v >= 1e6:  return f"${v/1e6:.0f}M"
        except Exception: pass
        return "N/A"

    def badge(txt, col):
        return f'<span style="background:{col};color:white;padding:3px 10px;border-radius:12px;font-size:11px;font-weight:700;white-space:nowrap">{txt}</span>'

    def trend(price, s50, s200):
        try:
            p, s5, s2 = float(price), float(s50), float(s200)
            if np.isnan(s5) or np.isnan(s2): return "N/A"
            if p > s5 > s2:  return '<span style="color:#1a7a4a;font-weight:700">Uptrend</span>'
            if p < s5 < s2:  return '<span style="color:#a31515;font-weight:700">Downtrend</span>'
            if s5 > s2:      return '<span style="color:#2e6ba8">Golden X</span>'
            return '<span style="color:#e8a020">Mixed</span>'
        except Exception: return "N/A"

    TH = """<th>Ticker</th><th>Sleeve</th><th>Price</th><th>Day</th>
            <th>1W</th><th>1M</th><th>3M</th><th>1Y Sim</th>
            <th>RSI</th><th>Trend</th>
            <th>Alloc</th><th>Sim Value</th><th>Gain/Loss</th>
            <th>Vol</th><th>Sharpe</th><th>Sortino</th><th>Max DD</th><th>Beta</th><th>Win%</th>
            <th>Mkt Cap</th><th>Fwd P/E</th><th>Rev Grth</th><th>Gross Mgn</th>
            <th>Analyst</th><th>Upside</th><th>Next Earn</th>
            <th>Signal</th><th>Key Drivers</th>"""

    def build_row(ticker, data):
        cfg = data.get("config", PORTFOLIO.get(ticker, {}))
        t   = data.get("technical", {})
        f   = data.get("fundamental", {})
        s   = data.get("signal", {})
        m   = data.get("metrics", {})
        sp  = (sim.get("positions") or {}).get(ticker, {})

        price   = t.get("price", 0) or 0
        dc      = t.get("day_change_pct", 0) or 0
        rsi_v   = t.get("rsi", float("nan"))
        rets    = t.get("returns", {})
        alloc   = cfg.get("target_pct", 0) * PORTFOLIO_CAPITAL
        sval    = sp.get("simulated_value")
        gl      = sp.get("gain_loss")
        beta_v  = sp.get("beta")
        sim_ret = sp.get("return_pct")

        upside = None
        tgt    = f.get("target_mean")
        if tgt and price:
            try: upside = round(((float(tgt) - float(price)) / float(price)) * 100, 1)
            except Exception: pass

        rsi_color = "#1a7a4a" if isinstance(rsi_v, float) and not np.isnan(rsi_v) and rsi_v < 35 \
                    else ("#a31515" if isinstance(rsi_v, float) and not np.isnan(rsi_v) and rsi_v > 65 else "#555")

        ar = (f.get("analyst_rec") or "n/a").replace("_", " ").upper()
        ar_col = "#1a7a4a" if "BUY" in ar else ("#a31515" if "SELL" in ar else "#888")

        ne = f.get("next_earnings", "N/A") or "N/A"
        ne_html = f'<span style="color:#e8a020;font-weight:600">{ne}</span>' if ne != "N/A" else "N/A"

        return f"""<tr>
          <td><strong style="color:#0d2b55">{ticker}</strong><br>
            <span style="font-size:10px;color:#aaa">{cfg.get('name','')}</span><br>
            <span style="font-size:9px;color:#ccc;font-style:italic">{cfg.get('theme','')}</span></td>
          <td><span style="font-size:10px;background:#e8f0f8;padding:1px 6px;border-radius:4px;color:#2e6ba8">{cfg.get('sleeve','')}</span></td>
          <td style="font-weight:bold">${fmt(price,'','',2)}</td>
          <td style="color:{pc(dc)};font-weight:bold">{fmt(dc,'','%',2,True)}</td>
          <td style="color:{pc(rets.get('1w'))}">{fmt(rets.get('1w'),'','%',1,True)}</td>
          <td style="color:{pc(rets.get('1m'))}">{fmt(rets.get('1m'),'','%',1,True)}</td>
          <td style="color:{pc(rets.get('3m'))}">{fmt(rets.get('3m'),'','%',1,True)}</td>
          <td style="color:{pc(sim_ret)};font-weight:bold">{fmt(sim_ret,'','%',1,True)}</td>
          <td style="font-weight:bold;color:{rsi_color}">{fmt(rsi_v,'','',1)}</td>
          <td>{trend(price, t.get('sma50'), t.get('sma200'))}</td>
          <td>{'N/A' if not alloc else f'£{alloc:,.0f}'}</td>
          <td style="color:{pc(sval - alloc if sval else None)}">{('£' + f'{sval:,.2f}') if sval else 'N/A'}</td>
          <td style="color:{pc(gl)};font-weight:bold">{('+' if gl and gl >= 0 else '') + ('£' + f'{gl:,.2f}' if gl else 'N/A')}</td>
          <td>{fmt(m.get('volatility_pct'),'','%',1)}</td>
          <td style="color:{'#1a7a4a' if m.get('sharpe_ratio') and float(m['sharpe_ratio'])>1 else '#a31515' if m.get('sharpe_ratio') and float(m['sharpe_ratio'])<0 else '#555'}">{fmt(m.get('sharpe_ratio'),'','',2)}</td>
          <td>{fmt(m.get('sortino_ratio'),'','',2)}</td>
          <td style="color:#a31515">{fmt(m.get('max_drawdown_pct'),'','%',1)}</td>
          <td>{fmt(beta_v,'','',2)}</td>
          <td>{fmt(m.get('win_rate_pct'),'','%',1)}</td>
          <td style="font-size:11px">{fmt_mc(f.get('market_cap'))}</td>
          <td>{fmt(f.get('pe_forward'),'','x',1)}</td>
          <td style="color:{pc(f.get('revenue_growth'))}">{fmt((f.get('revenue_growth') or 0)*100,'','%',0,True)}</td>
          <td>{fmt((f.get('gross_margins') or 0)*100,'','%',0)}</td>
          <td style="color:{ar_col};font-size:11px">{ar}</td>
          <td style="color:{pc(upside)};font-weight:bold">{fmt(upside,'','%',1,True) if upside is not None else 'N/A'}</td>
          <td>{ne_html}</td>
          <td>{badge(s.get('signal','N/A'), s.get('color','#888'))}</td>
          <td style="font-size:11px;color:#555;max-width:200px">{'<br>'.join('• '+r for r in s.get('reasons',[])[:3])}</td>
        </tr>"""

    held_rows  = "".join(build_row(t, d) for t, d in analysis.items() if PORTFOLIO.get(t, {}).get("status") == "HELD" and "error" not in d)
    watch_rows = "".join(build_row(t, d) for t, d in analysis.items() if PORTFOLIO.get(t, {}).get("status") == "WATCH" and "error" not in d)

    # KPIs
    pm       = sim.get("portfolio_metrics", {})
    tot_ret  = sim.get("total_return_pct", 0)
    bm_ret   = sim.get("benchmark_return_pct", 0)
    alpha    = sim.get("alpha_pct", 0)
    tot_val  = sim.get("total_simulated_gbp", PORTFOLIO_CAPITAL)
    gl_tot   = sim.get("total_gain_loss_gbp", 0)

    kpi_data = [
        ("Portfolio Value",    f"£{tot_val:,.2f}",                          f"Started £{PORTFOLIO_CAPITAL:,.0f}",       "#0d2b55"),
        ("Total Return",       f"{'+' if tot_ret>=0 else ''}{tot_ret:.2f}%",f"{sim.get('simulation_period','')}",      pc(tot_ret)),
        ("vs S&P 500 Alpha",   f"{'+' if alpha>=0 else ''}{alpha:.2f}%",    f"S&P returned {bm_ret:.2f}%",             pc(alpha)),
        ("Gain / Loss",        f"{'+' if gl_tot>=0 else ''}£{gl_tot:,.2f}","Simulated P&L",                            pc(gl_tot)),
        ("Sharpe Ratio",       str(pm.get("sharpe_ratio","N/A")),           "Risk-adjusted return (>1 = good)",        "#0d2b55"),
        ("Sortino Ratio",      str(pm.get("sortino_ratio","N/A")),          "Downside risk-adjusted return",           "#0d2b55"),
        ("Max Drawdown",       f"{pm.get('max_drawdown_pct','N/A')}%",      "Worst peak-to-trough loss",               "#a31515"),
        ("Annualised Vol",     f"{pm.get('volatility_pct','N/A')}%",        "Annualised std deviation",                "#0d2b55"),
    ]
    kpi_html = "".join(f'<div class="kpi"><div class="kpi-label">{l}</div><div class="kpi-value" style="color:{c}">{v}</div><div class="kpi-sub">{s}</div></div>'
                       for l, v, s, c in kpi_data)

    # Sector allocation
    sec = {}
    for ticker, d in PORTFOLIO.items():
        if d["status"] == "HELD" and d["target_pct"] > 0:
            sec[d["sector"]] = sec.get(d["sector"], 0) + d["target_pct"]
    sec_html = "".join(f'<tr><td>{k}</td><td style="font-weight:bold">{v*100:.0f}%</td><td>£{v*PORTFOLIO_CAPITAL:,.0f}</td>'
                       f'<td><div style="background:#2e6ba8;height:10px;width:{int(v*500)}px;border-radius:2px"></div></td></tr>'
                       for k, v in sorted(sec.items(), key=lambda x: -x[1]))

    # Monthly performance
    prev_v  = PORTFOLIO_CAPITAL * 0.90
    mo_html = ""
    for dt, val in sorted(sim.get("monthly_values", {}).items()):
        chg = val - prev_v
        mo_html += f'<tr><td>{dt[:7]}</td><td style="font-weight:bold">£{val:,.2f}</td>' \
                   f'<td style="color:{pc(chg)}">{"+£" if chg>=0 else "-£"}{abs(chg):,.2f}</td>' \
                   f'<td style="color:{pc(chg)}">{((val/prev_v)-1)*100:+.2f}%</td></tr>'
        prev_v = val

    # Action alerts
    alerts = ""
    for ticker, data in analysis.items():
        sig = data.get("signal", {}).get("signal", "")
        if sig in ("STRONG BUY", "SELL", "REDUCE"):
            st = PORTFOLIO.get(ticker, {}).get("status", "")
            alerts += f'<div class="alert"><strong>{ticker}</strong> ({PORTFOLIO.get(ticker,{}).get("name","")}) ' \
                      f'- {badge(sig, data["signal"]["color"])} ' \
                      f'{"- WATCHLIST: review for entry" if st=="WATCH" else "- HELD: action required"}</div>'
    if not alerts:
        alerts = '<div class="alert-ok">No urgent signals today - portfolio stable</div>'

    # Earnings
    earn_rows = ""
    for ticker, data in sorted(analysis.items()):
        ne = (data.get("fundamental") or {}).get("next_earnings", "N/A")
        if ne and ne != "N/A":
            st = PORTFOLIO.get(ticker, {}).get("status", "WATCH")
            earn_rows += f'<tr><td><strong>{ticker}</strong></td><td>{PORTFOLIO.get(ticker,{}).get("name","")}</td>' \
                         f'<td style="color:#e8a020;font-weight:bold">{ne}</td>' \
                         f'<td>{"Held" if st=="HELD" else "Watchlist"}</td></tr>'

    return f"""<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Murdoch Capital - Dashboard {report_date}</title>
<style>
* {{ box-sizing:border-box; margin:0; padding:0; }}
body {{ font-family:'Segoe UI',Arial,sans-serif; background:#eef2f7; color:#1a1a2e; font-size:13px; }}
.header {{ background:linear-gradient(135deg,#0d2b55,#1b3a6b 60%,#2e6ba8); color:white; padding:28px 40px; }}
.header h1 {{ font-size:26px; font-weight:800; letter-spacing:1px; }}
.header p {{ font-size:13px; opacity:.75; margin-top:5px; }}
.meta {{ background:#1b3a6b; color:#a8c4e0; font-size:12px; padding:7px 40px; display:flex; gap:24px; flex-wrap:wrap; }}
.container {{ max-width:1700px; margin:0 auto; padding:20px 40px; }}
h2 {{ font-size:17px; color:#0d2b55; margin:28px 0 10px; padding-bottom:7px; border-bottom:3px solid #2e6ba8; }}
.kpi-grid {{ display:grid; grid-template-columns:repeat(4,1fr); gap:14px; margin:12px 0 20px; }}
.kpi {{ background:white; border-radius:8px; padding:16px 20px; box-shadow:0 2px 8px rgba(0,0,0,.07); border-left:4px solid #2e6ba8; }}
.kpi-label {{ font-size:11px; color:#888; text-transform:uppercase; letter-spacing:.5px; }}
.kpi-value {{ font-size:22px; font-weight:800; margin-top:3px; }}
.kpi-sub {{ font-size:11px; color:#999; margin-top:3px; }}
.card {{ background:white; border-radius:8px; box-shadow:0 2px 8px rgba(0,0,0,.07); margin:12px 0; overflow:hidden; }}
.card-inner {{ padding:16px 20px; }}
.tw {{ overflow-x:auto; }}
table {{ width:100%; border-collapse:collapse; white-space:nowrap; }}
th {{ background:#0d2b55; color:white; padding:8px 10px; font-size:11px; text-align:left; position:sticky; top:0; }}
td {{ padding:7px 10px; font-size:12px; border-bottom:1px solid #e8edf2; vertical-align:top; }}
tr:hover td {{ background:#f5f8ff; }}
.alert {{ background:#fff3cd; border-left:4px solid #e8a020; padding:10px 14px; border-radius:4px; margin:6px 0; }}
.alert-ok {{ background:#d4edda; border-left:4px solid #1a7a4a; padding:10px 14px; border-radius:4px; margin:6px 0; color:#1a7a4a; font-weight:600; }}
.grid2 {{ display:grid; grid-template-columns:1fr 1fr; gap:16px; }}
.footer {{ background:#0d2b55; color:#a8c4e0; text-align:center; font-size:11px; padding:16px; margin-top:30px; }}
code {{ background:#f0f4f8; padding:2px 6px; border-radius:3px; font-size:11px; }}
</style></head><body>

<div class="header">
  <h1>MURDOCH CAPITAL - INVESTOR DASHBOARD</h1>
  <p>£5,000 Moderate Risk Portfolio &nbsp;|&nbsp; 10-Year Horizon (2026-2036) &nbsp;|&nbsp; {report_date}</p>
</div>
<div class="meta">
  <span>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}</span>
  <span>Data Source: {DATA_PROVIDER.upper()}</span>
  <span>Held: {sum(1 for v in PORTFOLIO.values() if v['status']=='HELD')} positions</span>
  <span>Watchlist: {sum(1 for v in PORTFOLIO.values() if v['status']=='WATCH')} names</span>
  <span>Backtest: {sim.get('simulation_period','N/A')}</span>
  <span style="color:#f0c040">NOT FINANCIAL ADVICE</span>
</div>

<div class="container">

<h2>Portfolio Performance</h2>
<div class="kpi-grid">{kpi_html}</div>

<h2>Action Alerts</h2>
<div class="card"><div class="card-inner">{alerts}</div></div>

<div class="grid2">
  <div>
    <h2>Sector Allocation</h2>
    <div class="card"><div class="tw"><table>
      <thead><tr><th>Sector</th><th>%</th><th>Capital</th><th>Weight</th></tr></thead>
      <tbody>{sec_html}</tbody>
    </table></div></div>
  </div>
  <div>
    <h2>Monthly Simulated Performance</h2>
    <div class="card"><div class="tw"><table>
      <thead><tr><th>Month</th><th>Value</th><th>Change</th><th>Return</th></tr></thead>
      <tbody>{mo_html if mo_html else '<tr><td colspan="4" style="color:#aaa;text-align:center">No monthly data - run with internet access</td></tr>'}</tbody>
    </table></div></div>
  </div>
</div>

<h2>Portfolio Holdings - Full Analysis</h2>
<p style="color:#888;margin-bottom:8px;font-size:11px">12-month simulated returns based on historical price data. Target allocations at inception (£{PORTFOLIO_CAPITAL:,.0f}).</p>
<div class="card"><div class="tw"><table>
  <thead><tr>{TH}</tr></thead>
  <tbody>{held_rows if held_rows else '<tr><td colspan="28" style="text-align:center;color:#aaa;padding:20px">No data - check internet connection and run again</td></tr>'}</tbody>
</table></div></div>

<h2>Watchlist - All Names Tracked Daily</h2>
<p style="color:#888;margin-bottom:8px;font-size:11px">BUY or STRONG BUY on any watchlist name triggers an entry review. All metrics identical to held positions.</p>
<div class="card"><div class="tw"><table>
  <thead><tr>{TH}</tr></thead>
  <tbody>{watch_rows if watch_rows else '<tr><td colspan="28" style="text-align:center;color:#aaa;padding:20px">No data - check internet connection and run again</td></tr>'}</tbody>
</table></div></div>

<h2>Upcoming Earnings Calendar</h2>
<div class="card"><div class="tw"><table>
  <thead><tr><th>Ticker</th><th>Company</th><th>Next Earnings</th><th>Status</th></tr></thead>
  <tbody>{earn_rows if earn_rows else '<tr><td colspan="4" style="color:#aaa;text-align:center;padding:16px">No earnings data available</td></tr>'}</tbody>
</table></div></div>

<h2>Interactive Brokers - Upgrade Path</h2>
<div class="card"><div class="card-inner" style="line-height:1.8;color:#444">
  Currently using <strong>Yahoo Finance</strong> (free, 15-min delayed). To upgrade to live real-time data via Interactive Brokers:<br>
  <ol style="margin:10px 0 0 20px">
    <li>Open an account at <strong>ibkr.co.uk</strong> (UK regulated)</li>
    <li>Enable <strong>TWS API</strong> in IB Account Management</li>
    <li>Install: <code>pip install ibapi</code></li>
    <li>Run IB Gateway locally on port 7497</li>
    <li>In this script, change: <code>DATA_PROVIDER = "yahoo"</code> to <code>DATA_PROVIDER = "ibkr"</code></li>
    <li>Implement the three stub methods in the <code>InteractiveBrokersProvider</code> class</li>
  </ol>
  <p style="margin-top:10px">IBKR provides: live tick data, real-time P&L, actual position quantities, order execution, and institutional fundamentals via Reuters/Refinitiv.</p>
</div></div>

</div>
<div class="footer">
  MURDOCH CAPITAL &nbsp;|&nbsp; Cameron Murdoch Portfolio Management &nbsp;|&nbsp; {report_date}<br>
  Simulated returns are hypothetical based on historical data. NOT regulated financial advice. Past performance does not guarantee future results.<br>
  Analytics Engine v2.0 &nbsp;|&nbsp; Provider: {DATA_PROVIDER.upper()} &nbsp;|&nbsp; Switch DATA_PROVIDER to "ibkr" for live data
</div>
</body></html>"""


# ─────────────────────────────────────────────────────────────────────────────
# DOCX REPORT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _bold_run(para, text, size_pt=11, color_hex=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16))
    return run

def generate_docx_report(analysis, sim, today, out_dir):
    """Generate a clean Word document daily report."""
    if not DOCX_AVAILABLE:
        print("  [DOCX] python-docx not installed - skipping Word report.")
        print("         Run: pip install python-docx")
        return

    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(title, "MURDOCH CAPITAL", size_pt=22, color_hex="1a3a5c")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Daily Portfolio Report  |  {today}")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = meta.add_run(f"Provider: {DATA_PROVIDER.upper()}  |  Capital: £{PORTFOLIO_CAPITAL:,.0f}  |  Horizon: 10 Years  |  Risk: Moderate")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Portfolio Performance KPIs ────────────────────────────────────────────
    h = doc.add_paragraph()
    _bold_run(h, "PORTFOLIO PERFORMANCE  (12-Month Simulation)", size_pt=13, color_hex="1a3a5c")
    h.paragraph_format.space_after = Pt(4)

    kpi_headers = ["Simulated Value", "Total Return", "vs S&P 500 (Alpha)", "Sharpe Ratio", "Max Drawdown"]
    pm = sim.get("portfolio_metrics", {})
    kpi_vals = [
        f"£{sim.get('total_simulated_gbp', PORTFOLIO_CAPITAL):,.2f}",
        f"{sim.get('total_return_pct', 0):+.2f}%",
        f"{sim.get('alpha_pct', 0):+.2f}%",
        str(pm.get("sharpe_ratio", "N/A")),
        f"{pm.get('max_drawdown_pct', 'N/A')}%"  if isinstance(pm.get('max_drawdown_pct'), (int, float)) else "N/A",
    ]

    kpi_table = doc.add_table(rows=2, cols=5)
    kpi_table.style = "Table Grid"
    for i, (hdr, val) in enumerate(zip(kpi_headers, kpi_vals)):
        hc = kpi_table.cell(0, i)
        vc = kpi_table.cell(1, i)
        _set_cell_bg(hc, "1a3a5c")
        _set_cell_bg(vc, "f0f4f8")
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(hp, hdr, size_pt=9, color_hex="ffffff")
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(val)
        vr.bold = True
        vr.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ── Action Signals ────────────────────────────────────────────────────────
    urgent = [(t, d["signal"]) for t, d in analysis.items()
              if d.get("signal", {}).get("signal") in ("STRONG BUY", "SELL", "REDUCE")]

    if urgent:
        h2 = doc.add_paragraph()
        _bold_run(h2, "ACTION SIGNALS", size_pt=13, color_hex="c0392b")
        h2.paragraph_format.space_after = Pt(4)

        act_table = doc.add_table(rows=1, cols=4)
        act_table.style = "Table Grid"
        act_headers = ["Ticker", "Signal", "Score", "Reason"]
        for i, hdr in enumerate(act_headers):
            c = act_table.cell(0, i)
            _set_cell_bg(c, "c0392b")
            p = c.paragraphs[0]
            _bold_run(p, hdr, size_pt=9, color_hex="ffffff")

        for ticker, sig in urgent:
            row = act_table.add_row()
            vals = [ticker, sig.get("signal", ""), str(sig.get("score", "")), sig.get("reason", "")]
            signal_name = sig.get("signal", "")
            bg = "fff3cd" if signal_name == "STRONG BUY" else "fde8e8" if signal_name in ("SELL","REDUCE") else "f8f9fa"
            for i, v in enumerate(vals):
                c = row.cells[i]
                _set_cell_bg(c, bg)
                p = c.paragraphs[0]
                r = p.add_run(v)
                r.font.size = Pt(10)
                if i == 1:
                    r.bold = True

        doc.add_paragraph()

    # ── Holdings table ────────────────────────────────────────────────────────
    def add_holdings_section(title_text, status_filter, header_color):
        rows = [(t, d) for t, d in analysis.items()
                if PORTFOLIO[t]["status"] == status_filter and "error" not in d]
        if not rows:
            return

        h3 = doc.add_paragraph()
        _bold_run(h3, title_text, size_pt=13, color_hex=header_color)
        h3.paragraph_format.space_after = Pt(4)

        cols = ["Ticker", "Name", "Sleeve", "Price", "RSI", "Signal", "1M Ret", "CAGR", "Sharpe", "Drawdown"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"

        # Header row
        for i, hdr in enumerate(cols):
            c = tbl.cell(0, i)
            _set_cell_bg(c, header_color)
            p = c.paragraphs[0]
            _bold_run(p, hdr, size_pt=8, color_hex="ffffff")

        # Data rows
        for ticker, d in rows:
            t_cfg = PORTFOLIO[ticker]
            tech  = d.get("technical", {})
            sig   = d.get("signal", {})
            met   = d.get("metrics", {})
            rets  = tech.get("returns", {})

            signal_name = sig.get("signal", "HOLD")
            bg = "e8f5e9" if "BUY" in signal_name else "fde8e8" if signal_name in ("SELL","REDUCE") else "fafafa"

            price_s  = f"${tech.get('price', 0):.2f}"
            rsi_val  = tech.get("rsi")
            rsi_s    = f"{rsi_val:.1f}" if isinstance(rsi_val, float) and not np.isnan(rsi_val) else "N/A"
            ret_1m   = rets.get("1m")
            ret_s    = f"{ret_1m:+.1f}%" if ret_1m is not None else "N/A"
            cagr_s   = f"{met.get('cagr_pct','N/A')}%" if isinstance(met.get('cagr_pct'), (int,float)) else "N/A"
            sharpe_s = str(met.get("sharpe_ratio", "N/A"))
            dd_s     = f"{met.get('max_drawdown_pct','N/A')}%" if isinstance(met.get('max_drawdown_pct'), (int,float)) else "N/A"

            row_vals = [ticker, t_cfg["name"][:22], t_cfg["sleeve"], price_s,
                        rsi_s, signal_name, ret_s, cagr_s, sharpe_s, dd_s]

            row = tbl.add_row()
            for i, v in enumerate(row_vals):
                c = row.cells[i]
                _set_cell_bg(c, bg)
                p = c.paragraphs[0]
                r = p.add_run(str(v))
                r.font.size = Pt(8.5)
                if i == 0:
                    r.bold = True

        doc.add_paragraph()

    add_holdings_section("HOLDINGS  (18 positions)", "HELD",  "1a3a5c")
    add_holdings_section("WATCHLIST  (10 names)",    "WATCH", "2d6a4f")

    # ── Monthly performance ───────────────────────────────────────────────────
    monthly = sim.get("monthly_values", {})
    if monthly:
        h4 = doc.add_paragraph()
        _bold_run(h4, "SIMULATED MONTHLY PERFORMANCE", size_pt=13, color_hex="1a3a5c")
        h4.paragraph_format.space_after = Pt(4)

        months_sorted = sorted(monthly.keys())
        mt = doc.add_table(rows=1, cols=len(months_sorted))
        mt.style = "Table Grid"
        for i, m in enumerate(months_sorted):
            c = mt.cell(0, i)
            _set_cell_bg(c, "1a3a5c")
            p = c.paragraphs[0]
            _bold_run(p, m, size_pt=7, color_hex="ffffff")
        row2 = mt.add_row()
        for i, m in enumerate(months_sorted):
            val = monthly[m]
            c = row2.cells[i]
            _set_cell_bg(c, "f0f4f8")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"£{val:,.0f}")
            r.font.size = Pt(8)

        doc.add_paragraph()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    disc = doc.add_paragraph()
    disc.paragraph_format.space_before = Pt(12)
    dr = disc.add_run(
        "DISCLAIMER: Simulated returns are hypothetical, based on historical data. "
        "This report is for informational purposes only and does not constitute regulated financial advice. "
        "Past performance does not guarantee future results. Murdoch Capital is not FCA regulated. "
        f"Analytics Engine v2.0  |  Provider: {DATA_PROVIDER.upper()}  |  Generated: {today}"
    )
    dr.font.size = Pt(8)
    dr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    dr.italic = True

    # ── Save ──────────────────────────────────────────────────────────────────
    docx_path = os.path.join(out_dir, f"portfolio_report_{today}.docx")
    doc.save(docx_path)
    print(f"  Saved: portfolio_report_{today}.docx")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    today    = date.today().isoformat()
    base_dir = os.path.dirname(os.path.abspath(__file__))
    out_dir  = os.path.join(base_dir, "Daily Reports", today)
    os.makedirs(out_dir, exist_ok=True)
    provider = get_provider()

    print(f"\n{'='*62}")
    print(f"  MURDOCH CAPITAL - ANALYTICS ENGINE v2.0")
    print(f"  {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  {DATA_PROVIDER.upper()}")
    print(f"{'='*62}")

    # 1. Backtest simulation
    sim = simulate_portfolio(provider, BACKTEST_MONTHS)
    if "error" in sim:
        sim = {"positions": {}, "portfolio_metrics": {}, "monthly_values": {},
               "total_return_pct": 0, "benchmark_return_pct": 0, "alpha_pct": 0,
               "total_simulated_gbp": PORTFOLIO_CAPITAL, "total_gain_loss_gbp": 0,
               "simulation_period": "N/A"}

    # 2. Analyse every ticker (held + watchlist)
    all_tickers = list(PORTFOLIO.keys())
    analysis    = {}
    end   = date.today()
    start = end - timedelta(days=400)

    print(f"\n  Analysing {len(all_tickers)} securities...\n")
    print(f"  {'Ticker':<10} {'Status':<10} {'Price':>8}  {'RSI':>6}  {'Signal':<14}  {'1M Ret':>8}")
    print(f"  {'-'*62}")

    for ticker in all_tickers:
        status = PORTFOLIO[ticker]["status"]
        try:
            hist   = provider.get_history(ticker, str(start), str(end))
            result = analyse_ticker(ticker, provider, hist)
            analysis[ticker] = result
            if "error" not in result:
                t = result["technical"]
                s = result["signal"]
                r = t.get("rsi", float("nan"))
                r_s = f"{r:.1f}" if isinstance(r, float) and not np.isnan(r) else " N/A"
                m1  = t.get("returns", {}).get("1m")
                m1s = f"{m1:+.1f}%" if m1 is not None else "  N/A"
                print(f"  {ticker:<10} {status:<10} {t.get('price',0):>8.2f}  {r_s:>6}  {s.get('signal','N/A'):<14}  {m1s:>8}")
            else:
                print(f"  {ticker:<10} {status:<10} {'ERROR - no data':>30}")
        except Exception as e:
            analysis[ticker] = {"error": str(e), "config": PORTFOLIO.get(ticker, {})}
            print(f"  {ticker:<10} {status:<10} {'ERROR':>8}  {'':>6}  {str(e)[:30]}")

    # 3. Generate dashboard
    print(f"\n  Generating dashboard...")
    html = generate_dashboard(analysis, sim, today)
    path = os.path.join(out_dir, f"portfolio_dashboard_{today}.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  Saved: portfolio_dashboard_{today}.html")

    # 3b. Generate Word report
    generate_docx_report(analysis, sim, today, out_dir)

    # 4. Save state
    state = {"date": today, "provider": DATA_PROVIDER,
             "sim_summary": {k: v for k, v in sim.items() if k not in ("positions","monthly_values")},
             "signals": {t: d.get("signal", {}) for t, d in analysis.items() if "signal" in d}}
    with open(os.path.join(out_dir, "portfolio_state.json"), "w") as f:
        json.dump(state, f, indent=2, default=str)

    # 5. Summary
    print(f"\n{'='*62}")
    print(f"  SUMMARY")
    print(f"  Simulated Value:  £{sim.get('total_simulated_gbp', PORTFOLIO_CAPITAL):>10,.2f}")
    print(f"  Total Return:      {sim.get('total_return_pct', 0):>+9.2f}%")
    print(f"  vs S&P 500:        {sim.get('alpha_pct', 0):>+9.2f}% alpha")
    pm = sim.get("portfolio_metrics", {})
    if pm:
        print(f"  Sharpe Ratio:      {str(pm.get('sharpe_ratio','N/A')):>10}")
        print(f"  Max Drawdown:      {str(pm.get('max_drawdown_pct','N/A')):>9}%")
    print(f"\n  Action signals:")
    acts = [(t, d["signal"]["signal"]) for t, d in analysis.items()
            if d.get("signal", {}).get("signal") in ("STRONG BUY", "SELL", "REDUCE")]
    for t, sig in acts:
        print(f"  ** {t:<10} {sig}")
    if not acts:
        print(f"  No urgent signals")
    print(f"\n  Files saved to: Daily Reports/{today}/")
    print(f"    portfolio_dashboard_{today}.html  (open in browser)")
    print(f"    portfolio_report_{today}.docx     (open in Word)")
    print(f"{'='*62}\n")


if __name__ == "__main__":
    main()
